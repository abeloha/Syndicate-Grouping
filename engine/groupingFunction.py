import operator
from itertools import groupby

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from datetime import datetime

reverse = False

class Student:
    def __init__(self, id, name, grade, gender, specialisation, foreigner=""):
        self.id = id
        self.name = name
        self.grade = grade
        self.gender = gender
        self.foreigner = foreigner
        self.specialisation = specialisation

    def __repr__(self):
        return repr((self.id, self.name, self.grade, self.gender, self.foreigner, self.specialisation))

def reverse_sort():
    global reverse
    reverse = not reverse
    return reverse

def perform_grouping(param_students_data, param_number_of_groups):
    students = []

    for d in param_students_data:
        id = d[0]
        name = d[3]
        country = d[5]
        gender = d[6]
        grade = d[7]
        specialty = d[8]
        students.append(Student(id,name,grade,gender,specialty,country))

    NUMBER_OF_GROUPS = param_number_of_groups
    groups = [[] for i in range(NUMBER_OF_GROUPS)]
    counter = 0

    gender_attr = operator.attrgetter("gender")
    grade_attr = operator.attrgetter("grade")
    natinal_attr = operator.attrgetter("foreigner")
    spe_attr = operator.attrgetter("specialisation")

    grouped_students = []
    # Group by national
    for national_key, national_vals in groupby(sorted(students, key=natinal_attr), natinal_attr):
        grouped_students.append(
            [
                [
                    list(grade_vals)
                    # Group by grade
                    for grade_key, grade_vals in groupby(sorted(gender_vals, key=grade_attr), grade_attr)
                ]
                # Group by gender
                for gender_key, gender_vals in groupby(sorted(national_vals, key=gender_attr), gender_attr)
            ]
        )

    for gender_group in grouped_students:
        for national_group in gender_group:
            for specialisation_group in national_group:
                for student in sorted(specialisation_group, key=spe_attr, reverse=reverse_sort()):
                    groups[counter].append(student)
                    counter = (counter + 1) % NUMBER_OF_GROUPS
    
    return groups

def print_group(student_groups):
    for i in range(len(student_groups)):
        print("Group {0}:".format(i + 1))
        for member in student_groups[i]:
            print(member)

def print_to_doc(data, filename, title, syntype): #1 = single service, 2 = tri-service
    
        if (not len(data)):
           return 0

        if (not filename):
            return 0


        group = ''

        document = Document()
        document.add_picture('resources\icon.png', width=Inches(1))
        last_par = document.paragraphs[-1]
        last_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_heading(title, 0)
        last_par = document.paragraphs[-1]
        last_par.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sn = 1
        for d in data:
            id = d[0]
            rank = d[2]
            name = d[3]
            p_no = d[4]
            country = d[5]
            gender = d[6]
            grade = d[7]
            specialty = d[8]
            remarks = d[9]

            if(syntype == 2):
                syn = d[14]
            else:
                syn = d[13]

            if (syn == 0):
                syn = "NONE"

            if(group != syn):                
                group = syn

                pr1 = document.add_paragraph('')
                document.add_heading('SYNDICATE ' + str(group), level=1)                
                pr2 = document.add_paragraph('')

                table = document.add_table(rows=0, cols=8)
                hdr_cells = table.add_row().cells
                hdr_cells[0].text = 'S/N'
                hdr_cells[1].text = 'Rank'
                hdr_cells[2].text = 'Name'
                hdr_cells[3].text = 'SVC No'
                hdr_cells[4].text = 'Corps'
                hdr_cells[5].text = 'Sex'
                hdr_cells[6].text = 'Country'
                hdr_cells[7].text = 'Grade'
                sn = 1


            row_cells = table.add_row().cells
            row_cells[0].text = str(sn)
            row_cells[1].text = str(rank)
            row_cells[2].text = str(name)
            row_cells[3].text = str(p_no)
            row_cells[4].text = str(specialty)
            row_cells[5].text = str(gender)
            row_cells[6].text = str(country)
            row_cells[7].text = str(grade)
            sn += 1
        

        date_created = datetime.now().strftime("%d-%m-%Y %H:%M")
        pr1 = document.add_paragraph('')
        pr2 = document.add_paragraph('')
        p = document.add_paragraph('Generated on ')
        p.add_run(date_created).bold = True

        try:
            document.save(str(filename) + '.docx')
            return 1
        except ValueError:
            pass
        
        return 0